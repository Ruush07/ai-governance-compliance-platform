"""Text extraction (Phase 2 implementation).

Concrete extractors for the formats the platform accepts:

* PDF  — PyMuPDF (``fitz``) for page text + layout, pdfplumber for table
  detection, with a pytesseract OCR fallback for scanned pages.
* DOCX — python-docx (paragraphs + tables).
* Images (png/jpg/tiff) — pytesseract OCR.

Heavy libraries are imported lazily *inside* the concrete extractors so
importing this module (and booting Django) never requires the optional
ingestion dependencies. The normalised :class:`ExtractedDocument` contract is
unchanged from the Phase 1 scaffold, so downstream code is unaffected.
"""
from __future__ import annotations

from abc import ABC, abstractmethod
from dataclasses import dataclass, field

# Heuristic: a page yielding fewer than this many characters of native text is
# treated as scanned/image-only and routed through OCR.
SCANNED_TEXT_THRESHOLD = 20


@dataclass
class ExtractedTable:
    page_number: int
    rows: list[list[str]] = field(default_factory=list)


@dataclass
class ExtractedPage:
    page_number: int
    text: str
    is_scanned: bool = False
    tables: list[ExtractedTable] = field(default_factory=list)


@dataclass
class ExtractedDocument:
    """Normalised output of any text extractor — the pipeline's ingestion
    contract. Page-level structure is preserved so evidence can cite pages
    (Rule 2: every quote requires a page citation)."""

    pages: list[ExtractedPage]
    metadata: dict = field(default_factory=dict)

    @property
    def page_count(self) -> int:
        return len(self.pages)

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)

    def build_page_map(self) -> list[dict]:
        """Compute global char-offset ranges per page over ``full_text``.

        Mirrors the join used by :pyattr:`full_text` (``"\\n\\n"`` separators)
        so offsets are exact and reversible to a page number.
        """
        page_map: list[dict] = []
        cursor = 0
        for i, page in enumerate(self.pages):
            start = cursor
            end = start + len(page.text)
            page_map.append({"page": page.page_number, "char_start": start, "char_end": end})
            # account for the "\n\n" separator between pages
            cursor = end + (2 if i < len(self.pages) - 1 else 0)
        return page_map


class TextExtractor(ABC):
    supported_extensions: set[str] = set()

    @abstractmethod
    def extract(self, file_path: str) -> ExtractedDocument:  # pragma: no cover
        ...


class PdfExtractor(TextExtractor):
    supported_extensions = {"pdf"}

    def extract(self, file_path: str) -> ExtractedDocument:
        import fitz  # PyMuPDF

        tables_by_page = self._extract_tables(file_path)
        pages: list[ExtractedPage] = []
        ocr_pages = 0
        with fitz.open(file_path) as doc:
            for index in range(doc.page_count):
                page = doc.load_page(index)
                text = page.get_text("text").strip()
                is_scanned = len(text) < SCANNED_TEXT_THRESHOLD
                if is_scanned:
                    ocr_text = self._ocr_page(page)
                    if ocr_text:
                        text = ocr_text
                        ocr_pages += 1
                pages.append(
                    ExtractedPage(
                        page_number=index + 1,
                        text=text,
                        is_scanned=is_scanned,
                        tables=tables_by_page.get(index + 1, []),
                    )
                )
        return ExtractedDocument(
            pages=pages,
            metadata={
                "extractor": "pymupdf",
                "ocr_pages": ocr_pages,
                "table_count": sum(len(t) for t in tables_by_page.values()),
            },
        )

    @staticmethod
    def _extract_tables(file_path: str) -> dict[int, list[ExtractedTable]]:
        """Best-effort table detection via pdfplumber (never fatal)."""
        result: dict[int, list[ExtractedTable]] = {}
        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    for raw in page.extract_tables() or []:
                        rows = [[(c or "") for c in row] for row in raw]
                        result.setdefault(i, []).append(ExtractedTable(page_number=i, rows=rows))
        except Exception:  # pragma: no cover - table detection is best-effort
            return result
        return result

    @staticmethod
    def _ocr_page(page) -> str:
        """Render a PDF page to an image and OCR it (best-effort)."""
        try:
            import fitz  # noqa: F401
            from .extract_image import ocr_image_bytes

            pix = page.get_pixmap(dpi=200)
            return ocr_image_bytes(pix.tobytes("png"))
        except Exception:  # pragma: no cover - OCR is a best-effort fallback
            return ""


class DocxExtractor(TextExtractor):
    supported_extensions = {"docx"}

    def extract(self, file_path: str) -> ExtractedDocument:
        import docx

        document = docx.Document(file_path)
        parts = [p.text for p in document.paragraphs if p.text.strip()]

        tables: list[ExtractedTable] = []
        for table in document.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(ExtractedTable(page_number=1, rows=rows))
            parts.extend(" | ".join(r) for r in rows)

        text = "\n".join(parts)
        return ExtractedDocument(
            pages=[ExtractedPage(page_number=1, text=text, tables=tables)],
            metadata={"extractor": "python-docx", "table_count": len(tables)},
        )


class ImageExtractor(TextExtractor):
    supported_extensions = {"png", "jpg", "jpeg", "tiff"}

    def extract(self, file_path: str) -> ExtractedDocument:
        from .extract_image import ocr_image_file

        text = ocr_image_file(file_path)
        return ExtractedDocument(
            pages=[ExtractedPage(page_number=1, text=text, is_scanned=True)],
            metadata={"extractor": "pytesseract", "ocr_pages": 1},
        )


_EXTRACTORS: list[TextExtractor] = [PdfExtractor(), DocxExtractor(), ImageExtractor()]


def extract_text(file_path: str, *, extension: str | None = None) -> ExtractedDocument:
    """Dispatch to the extractor matching ``extension`` (or the file suffix)."""
    if extension is None:
        extension = file_path.rsplit(".", 1)[-1].lower() if "." in file_path else ""
    extension = extension.lstrip(".").lower()
    for extractor in _EXTRACTORS:
        if extension in extractor.supported_extensions:
            return extractor.extract(file_path)
    raise ValueError(f"No text extractor registered for '.{extension}' files.")
